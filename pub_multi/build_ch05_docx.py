from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
DEFAULT_INPUT = ROOT / "doc" / "ch05_conclusion_formal_th.md"
DEFAULT_OUTPUT = ROOT / "out" / "บทที่5_สรุปผลการวิจัยและข้อเสนอแนะ.docx"
MAX_IMAGE_WIDTH = Cm(15.8)
CONTENT_FIRST_LINE_INDENT = Cm(1.27)


HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)$")
BULLET_RE = re.compile(r"^-\s+(.*)$")
TABLE_TITLE_RE = re.compile(r"^(ตารางที่\s+\d+\.\d+.*)$")
FIGURE_MARKER_RE = re.compile(r"^\[แทรกรูปที่\s+([0-9.]+)\s+ที่นี่(?:\s*ถ้าต้องการรูปปิดบท)?\]$")
FIGURE_REF_RE = re.compile(r"^\[(?:อ้างถึงรูปที่|เน้นอ้างถึงรูปที่)\s+(.*)\]$")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build Chapter 5 Word document from markdown and local figures.")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


def resolve_writable_output(path: Path) -> Path:
    if not path.exists():
        return path
    try:
        with open(path, "ab"):
            return path
    except PermissionError:
        stem = path.stem
        suffix = path.suffix
        parent = path.parent
        counter = 1
        while True:
            candidate = parent / f"{stem}_{counter}{suffix}"
            if not candidate.exists():
                return candidate
            try:
                with open(candidate, "ab"):
                    return candidate
            except PermissionError:
                counter += 1


def set_font(run_or_style, name: str, size_pt: float | None = None, bold: bool | None = None) -> None:
    font = run_or_style.font
    font.name = name
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    element = run_or_style._element if hasattr(run_or_style, "_element") else run_or_style.element
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.81)      # 1.5 inches
    section.right_margin = Cm(2.54)     # 1 inch
    section.top_margin = Cm(2.54)       # 1 inch
    section.bottom_margin = Cm(2.54)    # 1 inch
    section.header_distance = Cm(1.27)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_font(normal, "TH SarabunPSK", 16)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = CONTENT_FIRST_LINE_INDENT
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading1 = doc.styles["Heading 1"]
    set_font(heading1, "TH SarabunPSK", 20, True)
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(0)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.first_line_indent = Pt(0)

    heading2 = doc.styles["Heading 2"]
    set_font(heading2, "TH SarabunPSK", 18, True)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(0)
    heading2.paragraph_format.first_line_indent = Pt(0)

    heading3 = doc.styles["Heading 3"]
    set_font(heading3, "TH SarabunPSK", 16, True)
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(0)
    heading3.paragraph_format.first_line_indent = Pt(0)

    add_page_number_header(section)


def add_page_number_header(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    set_font(run, "TH SarabunPSK", 16)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, "Consolas", 10.5)
        else:
            subparts = re.split(r"(\*\*[^*]+\*\*)", part)
            for sub in subparts:
                if not sub:
                    continue
                if sub.startswith("**") and sub.endswith("**"):
                    run = paragraph.add_run(sub[2:-2])
                    set_font(run, "TH SarabunPSK", 16, True)
                else:
                    run = paragraph.add_run(sub)
                    set_font(run, "TH SarabunPSK", 16)


def is_alignment_row(row: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        stripped = line.strip().strip("|")
        row = [cell.strip() for cell in stripped.split("|")]
        rows.append(row)
    if len(rows) >= 2 and is_alignment_row(rows[1]):
        rows.pop(1)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (r_idx == 0 or c_idx > 0) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            if not text:
                continue
            run = p.add_run(text)
            set_font(run, "TH SarabunPSK", 15, bold=(r_idx == 0))
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    if level == 1:
        chapter_match = re.match(r"^(บทที่\s+\d+)\s+(.+)$", text)
        if chapter_match:
            p1 = doc.add_paragraph(style="Heading 1")
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.paragraph_format.first_line_indent = Pt(0)
            run1 = p1.add_run(chapter_match.group(1))
            set_font(run1, "TH SarabunPSK", 20, True)

            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(12)
            p2.paragraph_format.first_line_indent = Pt(0)
            run2 = p2.add_run(chapter_match.group(2))
            set_font(run2, "TH SarabunPSK", 24, True)
            return
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    add_inline_runs(p, text)


def add_paragraph(doc: Document, text: str, style: str | None = None, alignment: WD_ALIGN_PARAGRAPH | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    if style is None:
        p.paragraph_format.first_line_indent = CONTENT_FIRST_LINE_INDENT
        p.paragraph_format.line_spacing = 1.0
    add_inline_runs(p, text)


def add_table_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    match = re.match(r"^(ตารางที่\s+\d+\.\d+)\s+(.*)$", text)
    if match:
        run1 = p.add_run(match.group(1) + " ")
        set_font(run1, "TH SarabunPSK", 16, True)
        run2 = p.add_run(match.group(2))
        set_font(run2, "TH SarabunPSK", 16, False)
    else:
        run = p.add_run(text)
        set_font(run, "TH SarabunPSK", 16, True)


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        raise FileNotFoundError(f"Missing figure: {image_path}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=MAX_IMAGE_WIDTH)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Pt(0)
    match = re.match(r"^(รูปที่\s+\d+\.\d+)\s+(.*)$", caption)
    if match:
        run1 = cap.add_run(match.group(1) + " ")
        set_font(run1, "TH SarabunPSK", 15, True)
        run2 = cap.add_run(match.group(2))
        set_font(run2, "TH SarabunPSK", 15, False)
    else:
        run = cap.add_run(caption)
        set_font(run, "TH SarabunPSK", 15, False)
    doc.add_paragraph()


def consume_paragraph(lines: list[str], start: int) -> tuple[str, int]:
    parts = []
    i = start
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            break
        if (
            HEADING_RE.match(stripped)
            or stripped.startswith("|")
            or FIGURE_MARKER_RE.match(stripped)
            or FIGURE_REF_RE.match(stripped)
            or TABLE_TITLE_RE.match(stripped)
            or NUMBERED_RE.match(stripped)
            or BULLET_RE.match(stripped)
        ):
            break
        parts.append(stripped)
        i += 1
    return " ".join(parts), i


def build_doc(input_path: Path, output_path: Path) -> Path:
    text = input_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_document(doc)

    i = 0
    while i < len(lines):
        raw_line = lines[i].rstrip()
        stripped = raw_line.strip()
        if not stripped:
            i += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            add_heading(doc, heading_match.group(2).strip(), level)
            i += 1
            continue

        if TABLE_TITLE_RE.match(stripped):
            add_table_title(doc, stripped)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, parse_table(table_lines))
            continue

        figure_match = FIGURE_MARKER_RE.match(stripped)
        if figure_match:
            if i + 2 >= len(lines):
                raise ValueError(f"Incomplete figure marker near line {i + 1}")
            file_line = lines[i + 1].strip()
            caption_line = lines[i + 2].strip()
            if not file_line.startswith("ไฟล์รูป:") or not caption_line.startswith("คำบรรยาย:"):
                raise ValueError(f"Invalid figure marker block near line {i + 1}")
            image_path = Path(file_line.split(":", 1)[1].strip())
            caption = caption_line.split(":", 1)[1].strip()
            add_figure(doc, image_path, caption)
            i += 3
            continue

        ref_match = FIGURE_REF_RE.match(stripped)
        if ref_match:
            i += 1
            continue

        numbered_match = NUMBERED_RE.match(stripped)
        if numbered_match:
            add_paragraph(doc, numbered_match.group(1).strip(), style="List Number")
            i += 1
            continue

        bullet_match = BULLET_RE.match(stripped)
        if bullet_match:
            add_paragraph(doc, bullet_match.group(1).strip(), style="List Bullet")
            i += 1
            continue

        paragraph_text, next_idx = consume_paragraph(lines, i)
        if paragraph_text:
            add_paragraph(doc, paragraph_text)
            i = next_idx
        else:
            i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    final_output = resolve_writable_output(output_path)
    doc.save(final_output)
    return final_output


def main() -> None:
    args = parse_args()
    output_path = build_doc(args.input, args.output)
    print(output_path)


if __name__ == "__main__":
    main()
